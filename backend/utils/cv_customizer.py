from docx import Document
from docx.shared import Pt
import os
import re

# Predefined categories and skills
PREDEFINED_CATEGORIES = {
    "Cloud Platforms": [
        "AWS", "Azure", "Google Cloud Platform (GCP)", "EC2", "S3", "RDS", 
        "Lambda", "VPC", "AWS CloudFormation"
    ],
    "CI/CD Tools": [
        "Jenkins", "GitLab CI/CD", "GitHub Actions", "CircleCI", 
        "Bitbucket Pipelines", "SonarQube", "SonarCloud", "Nexus Repository"
    ],
    "Infrastructure as Code (IaC) Tools": [
        "Terraform", "Ansible", "CloudFormation"
    ],
    "Monitoring Tools": [
        "Prometheus", "Grafana", "ELK Stack (Elasticsearch, Logstash, Kibana)", "CloudWatch"
    ],
    "Scripting Languages": [
        "Python", "Bash", "PowerShell"
    ],
    "Containerization & Orchestration": [
        "Docker", "Kubernetes"
    ],
    "Version Control Tools": [
        "Git", "Bitbucket", "GitHub"
    ],
    "Operating Systems & Platforms": [
        "Linux", "Windows", "CentOS", "Ubuntu", "Red Hat Enterprise Linux (RHEL)"
    ],
    "Security Tools": [
        "HashiCorp Vault", "AWS IAM", "Azure AD", "SSL/TLS", 
        "SonarQube (Security Analysis)"
    ],
    "Database & Storage Tools": [
        "PostgreSQL", "MySQL", "MongoDB", "Redis", "DynamoDB", "ElastiCache", 
        "Cassandra", "Oracle DB"
    ],
    "Networking Tools": [
        "NGINX", "Apache", "iptables", "VLAN"
    ],
    "Artifact Repositories": [
        "Nexus Repository", "AWS CodeArtifact"
    ],
    "Build Tools": [
        "Maven", "Gradle", "Ant"
    ]
}

def sanitize_filename(filename):
    """
    Remove or replace invalid characters from the filename.
    """
    return re.sub(r'[<>:"/\\|?*$,]', '_', filename)

def extract_relevant_skills(job_description, predefined_categories):
    """
    Extract relevant skills from the job description for each predefined category.
    """
    matched_skills = {category: [] for category in predefined_categories}
    for category, skills in predefined_categories.items():
        for skill in skills:
            if skill.lower() in job_description.lower():
                matched_skills[category].append(skill)
    return matched_skills

def calculate_skill_match_percentage(matched_skills):
    """
    Calculate the skill match percentage for the job based on predefined categories.
    """
    total_skills = sum(len(skills) for skills in PREDEFINED_CATEGORIES.values())
    matched_skills_count = sum(len(skills) for skills in matched_skills.values())

    return (matched_skills_count / total_skills) * 100 if total_skills > 0 else 0

def customize_summary(original_summary, matched_skills):
    """
    Customize the professional summary by incorporating matched skills.
    """
    summary = original_summary
    for category, skills in matched_skills.items():
        if skills:
            skills_text = ", ".join(skills)
            summary += f" Skilled in {category.lower()} such as {skills_text}."
    return summary

def customize_cv(template_path, output_path, replacements):
    """
    Customize the CV template by replacing placeholders with specific content.
    """
    try:
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    # Save original formatting
                    is_bold = any(run.bold for run in paragraph.runs)
                    is_underline = any(run.underline for run in paragraph.runs)
                    font_size = next((run.font.size for run in paragraph.runs if run.font.size), Pt(11))

                    # Clear the paragraph text
                    paragraph.text = ""

                    # Add new text with formatting
                    new_run = paragraph.add_run(value)
                    new_run.font.size = font_size
                    new_run.bold = is_bold  # Retain bold formatting
                    new_run.underline = is_underline  # Retain underline formatting

        doc.save(output_path)
        print(f"File successfully written: {output_path}")
    except Exception as e:
        print(f"Error customizing CV: {e}")

def generate_cv(job, template_type, output_dir):
    """
    Generate a customized CV based on the job description and template type.
    Save the CV locally and return the path as a string.
    """
    root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    template_file = os.path.join(root_dir, "templates", f"{template_type}_cv_template.docx")

    if not os.path.exists(template_file):
        raise FileNotFoundError(f"Template file not found: {template_file}")

    sanitized_title = sanitize_filename(job['title'])
    output_file = os.path.join(output_dir, f"Joel_Maguranye_{sanitized_title}_CV.docx")

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Extract matched skills
    description = job.get("description", "N/A")
    matched_skills = extract_relevant_skills(description, PREDEFINED_CATEGORIES)

    # Prepare the summary and replacements
    original_summary = "DevOps Engineer with over 5 years of experience in cloud infrastructure, specializing in CI/CD pipelines and infrastructure automation."
    customized_summary = customize_summary(original_summary, matched_skills)

    replacements = {
        "{job_title}": job.get("title", "N/A"),
        "{skills}": ", ".join([", ".join(skills) for skills in matched_skills.values() if skills]),
        "{professional_summary}": customized_summary,
        "{job_duties}": description,
        "{company}": job.get("company", "N/A"),
        "{location}": job.get("location", "N/A"),
    }

    # Customize the CV
    customize_cv(template_file, output_file, replacements)

    print(f"File successfully written: {output_file}")
    return output_file  # Ensure only the file path is returned
